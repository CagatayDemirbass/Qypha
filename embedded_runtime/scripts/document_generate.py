import json
import os
import re
import sys
import tempfile
import zipfile
from typing import Any, Dict, List, Optional


def error(message: str, error_type: str = "runtime") -> None:
    print(
        json.dumps(
            {
                "ok": False,
                "error": message,
                "errorType": error_type,
            },
            ensure_ascii=False,
        )
    )
    raise SystemExit(1)


def read_spec() -> Dict[str, Any]:
    try:
        payload = json.load(sys.stdin)
    except Exception as exc:
        error(f"Failed to parse JSON input: {exc}", "input")
    if not isinstance(payload, dict):
        error("Document generation input must be a JSON object.", "input")
    return payload


def normalize_text(value: Any) -> Optional[str]:
    if value is None:
        return None
    if not isinstance(value, str):
        return None
    trimmed = value.strip()
    return trimmed or None


def normalize_string_list(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    result: List[str] = []
    for entry in value:
        if not isinstance(entry, str):
            continue
        trimmed = entry.strip()
        if trimmed:
            result.append(trimmed)
    return result


def normalize_sections(value: Any) -> List[Dict[str, Any]]:
    if not isinstance(value, list):
        return []
    sections: List[Dict[str, Any]] = []
    for entry in value:
        if not isinstance(entry, dict):
            continue
        heading = normalize_text(entry.get("heading"))
        text = normalize_text(entry.get("text"))
        paragraphs = normalize_string_list(entry.get("paragraphs"))
        bullets = normalize_string_list(entry.get("bullets"))
        if heading or text or paragraphs or bullets:
            sections.append(
                {
                    "heading": heading,
                    "text": text,
                    "paragraphs": paragraphs,
                    "bullets": bullets,
                }
            )
    return sections


def normalize_rows(value: Any) -> List[List[Any]]:
    if not isinstance(value, list):
        return []
    rows: List[List[Any]] = []
    for row in value:
        if isinstance(row, list):
            rows.append(
                [
                    cell
                    if isinstance(cell, (str, int, float, bool)) or cell is None
                    else str(cell)
                    for cell in row
                ]
            )
        else:
            rows.append([str(row)])
    return rows


def normalize_sheets(value: Any) -> List[Dict[str, Any]]:
    if not isinstance(value, list):
        return []
    sheets: List[Dict[str, Any]] = []
    for index, entry in enumerate(value):
        if not isinstance(entry, dict):
            continue
        raw_name = normalize_text(entry.get("name")) or f"Sheet{index + 1}"
        rows = normalize_rows(entry.get("rows"))
        charts = normalize_charts(entry.get("charts"))
        sheets.append({"name": raw_name[:31], "rows": rows, "charts": charts})
    return sheets


def normalize_positive_int(value: Any) -> Optional[int]:
    if isinstance(value, bool):
        return None
    if isinstance(value, int):
        return value if value >= 1 else None
    if isinstance(value, float) and value.is_integer():
        integer_value = int(value)
        return integer_value if integer_value >= 1 else None
    if isinstance(value, str):
        trimmed = value.strip()
        if trimmed.isdigit():
            integer_value = int(trimmed)
            return integer_value if integer_value >= 1 else None
    return None


def normalize_chart_type(value: Any) -> str:
    normalized = normalize_text(value)
    if normalized is None:
        return "column"
    lowered = normalized.lower().replace("-", " ").replace("_", " ")
    if lowered in {"column", "col", "column chart", "column plot", "vertical bar", "vertical bar chart"}:
        return "column"
    if lowered in {"bar", "bar chart", "bar plot", "horizontal bar", "horizontal bar chart"}:
        return "bar"
    return "column"


def normalize_chart_color(value: Any) -> Optional[str]:
    normalized = normalize_text(value)
    if normalized is None:
        return None
    candidate = normalized.lstrip("#").upper()
    if re.fullmatch(r"[0-9A-F]{6}", candidate):
        return candidate
    return None


def normalize_chart_colors(value: Any) -> List[str]:
    if not isinstance(value, list):
        return []
    colors: List[str] = []
    for entry in value:
        color = normalize_chart_color(entry)
        if color:
            colors.append(color)
    return colors


def normalize_charts(value: Any) -> List[Dict[str, Any]]:
    if not isinstance(value, list):
        return []
    charts: List[Dict[str, Any]] = []
    for entry in value:
        if not isinstance(entry, dict):
            continue
        chart_type = normalize_chart_type(entry.get("type"))
        title = normalize_text(entry.get("title"))
        anchor = normalize_text(entry.get("anchor")) or "D2"
        header_row = normalize_positive_int(entry.get("headerRow"))
        data_start_row = normalize_positive_int(entry.get("dataStartRow"))
        data_end_row = normalize_positive_int(entry.get("dataEndRow"))
        data_start_column = normalize_positive_int(entry.get("dataStartColumn"))
        data_end_column = normalize_positive_int(entry.get("dataEndColumn"))
        categories_column = normalize_positive_int(entry.get("categoriesColumn"))
        categories_start_row = normalize_positive_int(entry.get("categoriesStartRow"))
        categories_end_row = normalize_positive_int(entry.get("categoriesEndRow"))
        style = normalize_positive_int(entry.get("style"))
        x_axis_title = normalize_text(entry.get("xAxisTitle"))
        y_axis_title = normalize_text(entry.get("yAxisTitle"))
        series_colors = normalize_chart_colors(entry.get("seriesColors"))
        charts.append(
            {
                "type": chart_type,
                "title": title,
                "anchor": anchor,
                "headerRow": header_row,
                "dataStartRow": data_start_row,
                "dataEndRow": data_end_row,
                "dataStartColumn": data_start_column,
                "dataEndColumn": data_end_column,
                "categoriesColumn": categories_column,
                "categoriesStartRow": categories_start_row,
                "categoriesEndRow": categories_end_row,
                "style": style,
                "xAxisTitle": x_axis_title,
                "yAxisTitle": y_axis_title,
                "seriesColors": series_colors,
            }
        )
    return charts


INTEGER_LITERAL_RE = re.compile(r"^[+-]?\d+$")
FLOAT_LITERAL_RE = re.compile(
    r"^[+-]?(?:(?:\d+\.\d*)|(?:\d*\.\d+)|(?:\d+(?:[eE][+-]?\d+))|(?:\d+\.\d*[eE][+-]?\d+)|(?:\d*\.\d+[eE][+-]?\d+))$"
)


def coerce_chart_numeric_value(value: Any) -> Any:
    if not isinstance(value, str):
        return value
    trimmed = value.strip()
    if not trimmed:
        return value
    if INTEGER_LITERAL_RE.fullmatch(trimmed):
        signless = trimmed.lstrip("+-")
        if len(signless) > 1 and signless.startswith("0"):
            return value
        try:
            return int(trimmed)
        except ValueError:
            return value
    if FLOAT_LITERAL_RE.fullmatch(trimmed):
        signless = trimmed.lstrip("+-")
        if len(signless) > 1 and signless.startswith("0") and signless[1].isdigit():
            return value
        try:
            return float(trimmed)
        except ValueError:
            return value
    return value


def coerce_chart_data_cells_to_numbers(
    worksheet: Any,
    data_start_column: int,
    data_end_column: int,
    data_start_row: int,
    data_end_row: int,
) -> None:
    for row_index in range(data_start_row, data_end_row + 1):
        for column_index in range(data_start_column, data_end_column + 1):
            cell = worksheet.cell(row=row_index, column=column_index)
            coerced = coerce_chart_numeric_value(cell.value)
            if coerced is not cell.value:
                cell.value = coerced


def require_content(format_name: str, content: Optional[str], sections: List[Dict[str, Any]]) -> None:
    if content or sections:
        return
    error(f"{format_name} documents require content or sections.", "input")


def require_sheets(sheets: List[Dict[str, Any]]) -> None:
    if sheets:
        return
    error("xlsx documents require at least one sheet definition.", "input")


def ensure_parent_dir(target_path: str) -> None:
    parent = os.path.dirname(target_path) or "."
    os.makedirs(parent, exist_ok=True)


def allocate_temp_path(target_path: str) -> str:
    parent = os.path.dirname(target_path) or "."
    basename = os.path.basename(target_path)
    fd, temp_path = tempfile.mkstemp(prefix=f".{basename}.", suffix=".tmp", dir=parent)
    os.close(fd)
    return temp_path


def finalize_output(temp_path: str, target_path: str, overwrite: bool) -> int:
    if not overwrite and os.path.exists(target_path):
        try:
            os.unlink(temp_path)
        except FileNotFoundError:
            pass
        error(f"Target already exists: {target_path}", "input")
    os.replace(temp_path, target_path)
    return os.path.getsize(target_path)


def verify_pdf(path_value: str) -> None:
    with open(path_value, "rb") as handle:
        header = handle.read(5)
    if not header.startswith(b"%PDF"):
        error(f"Generated file is not a valid PDF: {path_value}")


def verify_zip_member(path_value: str, member: str) -> None:
    try:
        with zipfile.ZipFile(path_value, "r") as archive:
            if member not in archive.namelist():
                error(f"Generated archive is missing {member}: {path_value}")
    except zipfile.BadZipFile:
        error(f"Generated file is not a valid Office archive: {path_value}")


def add_paragraph_lines(story: List[Any], styles: Any, text: str) -> None:
    from reportlab.platypus import Paragraph, Spacer
    from xml.sax.saxutils import escape

    for block in text.split("\n\n"):
        trimmed = block.strip()
        if not trimmed:
            continue
        safe = escape(trimmed).replace("\n", "<br/>")
        story.append(Paragraph(safe, styles["BodyText"]))
        story.append(Spacer(1, 8))


def generate_pdf(
    target_path: str,
    title: Optional[str],
    content: Optional[str],
    sections: List[Dict[str, Any]],
) -> None:
    try:
        from reportlab.lib.pagesizes import LETTER
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.platypus import ListFlowable, ListItem, Paragraph, SimpleDocTemplate, Spacer
        from xml.sax.saxutils import escape
    except ImportError:
        error("reportlab is not installed; cannot generate PDF files.", "dependency")

    styles = getSampleStyleSheet()
    story: List[Any] = []
    if title:
        story.append(Paragraph(escape(title), styles["Title"]))
        story.append(Spacer(1, 16))
    if content:
        add_paragraph_lines(story, styles, content)
    for section in sections:
        heading = section.get("heading")
        if heading:
            story.append(Paragraph(escape(heading), styles["Heading2"]))
            story.append(Spacer(1, 8))
        text = section.get("text")
        if text:
            add_paragraph_lines(story, styles, text)
        for paragraph in section.get("paragraphs", []):
            add_paragraph_lines(story, styles, paragraph)
        bullets = section.get("bullets", [])
        if bullets:
            items = [
                ListItem(Paragraph(escape(bullet).replace("\n", "<br/>"), styles["BodyText"]))
                for bullet in bullets
            ]
            story.append(ListFlowable(items, bulletType="bullet"))
            story.append(Spacer(1, 8))
    if not story:
        error("PDF generation produced no content.", "input")
    doc = SimpleDocTemplate(target_path, pagesize=LETTER)
    doc.build(story)


def add_docx_text(document: Any, text: str) -> None:
    for block in text.split("\n\n"):
        trimmed = block.strip()
        if trimmed:
            document.add_paragraph(trimmed)


def generate_docx(
    target_path: str,
    title: Optional[str],
    content: Optional[str],
    sections: List[Dict[str, Any]],
) -> None:
    try:
        from docx import Document
    except ImportError:
        error("python-docx is not installed; cannot generate DOCX files.", "dependency")

    document = Document()
    if title:
        document.add_heading(title, level=0)
    if content:
        add_docx_text(document, content)
    for section in sections:
        heading = section.get("heading")
        if heading:
            document.add_heading(heading, level=1)
        text = section.get("text")
        if text:
            add_docx_text(document, text)
        for paragraph in section.get("paragraphs", []):
            add_docx_text(document, paragraph)
        for bullet in section.get("bullets", []):
            document.add_paragraph(bullet, style="List Bullet")
    document.save(target_path)


def generate_xlsx(target_path: str, title: Optional[str], sheets: List[Dict[str, Any]]) -> None:
    try:
        from openpyxl import Workbook
        from openpyxl.chart import BarChart, Reference
        from openpyxl.styles import Font
    except ImportError:
        error("openpyxl is not installed; cannot generate XLSX files.", "dependency")

    workbook = Workbook()
    first_sheet = True
    header_font = Font(bold=True)
    if title:
        workbook.properties.title = title
    for sheet_spec in sheets:
        worksheet = workbook.active if first_sheet else workbook.create_sheet()
        first_sheet = False
        worksheet.title = sheet_spec["name"]
        rows = sheet_spec.get("rows", [])
        for row_index, row in enumerate(rows, start=1):
            worksheet.append(row)
            if row_index == 1:
                for cell in worksheet[row_index]:
                    cell.font = header_font
        if rows:
            worksheet.freeze_panes = "A2"
        max_row = worksheet.max_row
        max_column = worksheet.max_column
        for chart_spec in sheet_spec.get("charts", []):
            if max_row <= 0 or max_column <= 0:
                continue
            data_start_column = chart_spec.get("dataStartColumn") or 1
            data_end_column = chart_spec.get("dataEndColumn") or max_column
            header_row = chart_spec.get("headerRow")
            if header_row is None:
                header_row = 1 if max_row >= 2 else None
            data_start_row = chart_spec.get("dataStartRow")
            if data_start_row is None:
                data_start_row = (header_row + 1) if header_row else 1
            data_end_row = chart_spec.get("dataEndRow") or max_row
            if (
                data_start_column > data_end_column
                or data_start_row > data_end_row
                or data_start_column > max_column
                or data_start_row > max_row
            ):
                continue

            coerce_chart_data_cells_to_numbers(
                worksheet,
                data_start_column=data_start_column,
                data_end_column=data_end_column,
                data_start_row=data_start_row,
                data_end_row=min(data_end_row, max_row),
            )

            chart = BarChart()
            chart.type = "bar" if chart_spec.get("type") == "bar" else "col"
            chart.style = chart_spec.get("style") or 10
            if chart_spec.get("title"):
                chart.title = chart_spec["title"]
            if chart_spec.get("xAxisTitle"):
                chart.x_axis.title = chart_spec["xAxisTitle"]
            if chart_spec.get("yAxisTitle"):
                chart.y_axis.title = chart_spec["yAxisTitle"]

            for column_index in range(data_start_column, data_end_column + 1):
                reference_start_row = header_row if header_row else data_start_row
                data_reference = Reference(
                    worksheet,
                    min_col=column_index,
                    max_col=column_index,
                    min_row=reference_start_row,
                    max_row=data_end_row,
                )
                chart.add_data(data_reference, titles_from_data=header_row is not None)

            categories_column = chart_spec.get("categoriesColumn")
            if categories_column is not None and categories_column <= max_column:
                categories_start_row = chart_spec.get("categoriesStartRow") or data_start_row
                categories_end_row = chart_spec.get("categoriesEndRow") or data_end_row
                if categories_start_row <= categories_end_row <= max_row:
                    categories_reference = Reference(
                        worksheet,
                        min_col=categories_column,
                        max_col=categories_column,
                        min_row=categories_start_row,
                        max_row=categories_end_row,
                    )
                    chart.set_categories(categories_reference)

            for series_index, color in enumerate(chart_spec.get("seriesColors", [])):
                if series_index >= len(chart.ser):
                    break
                graphical_props = chart.ser[series_index].graphicalProperties
                graphical_props.solidFill = color
                graphical_props.line.solidFill = color

            worksheet.add_chart(chart, chart_spec.get("anchor") or "D2")
    if first_sheet:
        workbook.active.title = "Sheet1"
    workbook.save(target_path)


def main() -> None:
    spec = read_spec()
    format_name = normalize_text(spec.get("format"))
    target_path = normalize_text(spec.get("path"))
    title = normalize_text(spec.get("title"))
    content = normalize_text(spec.get("content"))
    sections = normalize_sections(spec.get("sections"))
    sheets = normalize_sheets(spec.get("sheets"))
    overwrite = bool(spec.get("overwrite"))

    if format_name not in {"pdf", "docx", "xlsx"}:
        error("format must be one of pdf, docx, or xlsx.", "input")
    if not target_path:
        error("path is required.", "input")

    if format_name in {"pdf", "docx"}:
        require_content(format_name, content, sections)
    if format_name == "xlsx":
        require_sheets(sheets)

    ensure_parent_dir(target_path)
    temp_path = allocate_temp_path(target_path)
    try:
        if format_name == "pdf":
            generate_pdf(temp_path, title, content, sections)
            verify_pdf(temp_path)
        elif format_name == "docx":
            generate_docx(temp_path, title, content, sections)
            verify_zip_member(temp_path, "word/document.xml")
        else:
            generate_xlsx(temp_path, title, sheets)
            verify_zip_member(temp_path, "xl/workbook.xml")

        size_bytes = finalize_output(temp_path, target_path, overwrite)
    finally:
        if os.path.exists(temp_path):
            try:
                os.unlink(temp_path)
            except FileNotFoundError:
                pass

    print(
        json.dumps(
            {
                "ok": True,
                "format": format_name,
                "path": target_path,
                "sizeBytes": size_bytes,
                "title": title,
            },
            ensure_ascii=False,
        )
    )


if __name__ == "__main__":
    main()
